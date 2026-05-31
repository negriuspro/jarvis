"""
Unified file processor — uses Gemini Vision for images/video/docs,
pdfplumber for PDF text extraction, pandas for data files.
No ffmpeg required.
"""
import logging
import mimetypes
import os
from pathlib import Path

log = logging.getLogger("daniel.files")

_GEMINI_KEY = lambda: os.environ.get("GEMINI_API_KEY", "")
_GROQ_KEY = lambda: os.environ.get("GROQ_API_KEY", "")


# ─── Entry point ─────────────────────────────────────────────────────────────

def process_file(file_path: str, instruction: str = "") -> str:
    path = Path(file_path)
    if not path.exists():
        return f"Archivo no encontrado: {file_path}"

    ext = path.suffix.lower()
    mime, _ = mimetypes.guess_type(str(path))
    mime = mime or ""

    if ext in (".jpg", ".jpeg", ".png", ".gif", ".webp", ".bmp") or mime.startswith("image/"):
        return _process_image(path, instruction)

    if ext == ".pdf":
        return _process_pdf(path, instruction)

    if ext in (".mp4", ".mov", ".avi", ".mkv", ".webm") or mime.startswith("video/"):
        return _process_video(path, instruction)

    if ext in (".mp3", ".wav", ".ogg", ".m4a", ".flac") or mime.startswith("audio/"):
        return _process_audio(path, instruction)

    if ext in (".csv", ".xlsx", ".xls"):
        return _process_data(path, instruction)

    if ext in (".py", ".js", ".ts", ".java", ".cs", ".cpp", ".c", ".go", ".rs"):
        return _process_code(path, instruction)

    if ext in (".txt", ".md", ".docx", ".doc"):
        return _process_document(path, instruction)

    # Try generic Gemini analysis for unknown types
    try:
        content = path.read_text(encoding="utf-8", errors="ignore")[:8000]
        return _gemini_text(f"{instruction}\n\nContenido:\n{content}" if instruction else content)
    except Exception as e:
        return f"No sé cómo procesar este tipo de archivo ({ext}): {e}"


# ─── Image ────────────────────────────────────────────────────────────────────

def _process_image(path: Path, instruction: str) -> str:
    """Describe, OCR, resize, or compress using Gemini Vision + Pillow."""
    inst_lower = instruction.lower()

    # Resize / compress / convert — no AI needed
    if any(k in inst_lower for k in ("redimensiona", "resize", "comprime", "compress", "convierte", "convert")):
        return _image_transform(path, instruction)

    # OCR / describe via Gemini Vision
    prompt = instruction or "Describe esta imagen en detalle en español."
    if any(k in inst_lower for k in ("texto", "text", "ocr", "extrae", "extract", "lee")):
        prompt = "Extrae y transcribe todo el texto visible en esta imagen, exactamente como aparece."

    return _gemini_vision(path, prompt)


def _image_transform(path: Path, instruction: str) -> str:
    try:
        from PIL import Image
        img = Image.open(path)
        inst_lower = instruction.lower()
        original_size = path.stat().st_size

        # Resize
        if "redimensiona" in inst_lower or "resize" in inst_lower:
            import re
            nums = re.findall(r'\d+', instruction)
            if len(nums) >= 2:
                w, h = int(nums[0]), int(nums[1])
                img = img.resize((w, h), Image.LANCZOS)

        # Convert format
        out_format = img.format or "JPEG"
        out_ext = path.suffix
        if "png" in inst_lower:
            out_format, out_ext = "PNG", ".png"
        elif "jpg" in inst_lower or "jpeg" in inst_lower:
            out_format, out_ext = "JPEG", ".jpg"
        elif "webp" in inst_lower:
            out_format, out_ext = "WEBP", ".webp"

        out_path = path.with_name(path.stem + "_procesada" + out_ext)
        save_kwargs: dict = {"format": out_format}

        # Compress
        if "comprime" in inst_lower or "compress" in inst_lower:
            import re
            quality_match = re.search(r'(\d+)\s*%', instruction)
            quality = int(quality_match.group(1)) if quality_match else 75
            if out_format in ("JPEG", "WEBP"):
                save_kwargs["quality"] = quality
                save_kwargs["optimize"] = True

        if out_format == "JPEG" and img.mode in ("RGBA", "P"):
            img = img.convert("RGB")

        img.save(str(out_path), **save_kwargs)
        new_size = out_path.stat().st_size
        reduction = round((1 - new_size / original_size) * 100, 1)
        return (
            f"Imagen procesada guardada en: {out_path.name}\n"
            f"Tamaño: {_fmt_bytes(original_size)} → {_fmt_bytes(new_size)} "
            f"({'−' if reduction > 0 else '+'}{abs(reduction)}%)"
        )
    except Exception as e:
        return f"Error procesando imagen: {e}"


# ─── PDF ─────────────────────────────────────────────────────────────────────

def _process_pdf(path: Path, instruction: str) -> str:
    try:
        import pdfplumber
    except ImportError:
        return "pdfplumber no instalado. Ejecuta: pip install pdfplumber"

    inst_lower = instruction.lower()

    try:
        with pdfplumber.open(str(path)) as pdf:
            pages = pdf.pages
            total = len(pages)
            # Extract text from first 20 pages max
            text_parts = []
            for page in pages[:20]:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
            text = "\n\n".join(text_parts)
    except Exception as e:
        return f"No pude leer el PDF: {e}"

    if not text.strip():
        return "El PDF no tiene texto extraíble (puede ser un PDF escaneado). Intenta con un archivo de imagen."

    # Convert to Word
    if "word" in inst_lower or "docx" in inst_lower or "convierte" in inst_lower:
        return _pdf_to_word(path, text)

    # Summarize or analyze via Gemini
    truncated = text[:12000]
    prompt = instruction or f"Resume el siguiente documento PDF ({total} páginas) en español, destacando los puntos principales:"
    return _gemini_text(f"{prompt}\n\n{truncated}")


def _pdf_to_word(path: Path, text: str) -> str:
    try:
        from docx import Document
        doc = Document()
        doc.add_heading(path.stem, 0)
        for para in text.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
        out = path.with_suffix(".docx")
        doc.save(str(out))
        return f"PDF convertido a Word: {out.name}"
    except ImportError:
        return "python-docx no instalado. Ejecuta: pip install python-docx"
    except Exception as e:
        return f"Error convirtiendo PDF a Word: {e}"


# ─── Video ────────────────────────────────────────────────────────────────────

def _process_video(path: Path, instruction: str) -> str:
    """Transcribe video using Gemini File API (no ffmpeg needed)."""
    if not _GEMINI_KEY():
        return "GEMINI_API_KEY no configurada."
    if path.stat().st_size > 200 * 1024 * 1024:
        return "El video supera los 200MB. Usa un archivo más pequeño."

    try:
        import google.generativeai as genai
        genai.configure(api_key=_GEMINI_KEY())

        log.info("Subiendo video a Gemini File API: %s", path.name)
        video_file = genai.upload_file(str(path))

        # Wait for processing
        import time
        while video_file.state.name == "PROCESSING":
            time.sleep(3)
            video_file = genai.get_file(video_file.name)

        if video_file.state.name == "FAILED":
            return "Gemini no pudo procesar el video."

        model = genai.GenerativeModel("gemini-2.0-flash")
        prompt = instruction or "Transcribe todo el audio hablado en este video en español, con marcas de tiempo aproximadas."
        response = model.generate_content([video_file, prompt])
        genai.delete_file(video_file.name)
        return response.text

    except Exception as e:
        return f"Error procesando video: {e}"


# ─── Audio ────────────────────────────────────────────────────────────────────

def _process_audio(path: Path, instruction: str) -> str:
    """Transcribe audio using Groq Whisper."""
    if not _GROQ_KEY():
        return "GROQ_API_KEY no configurada."
    try:
        from groq import Groq
        client = Groq(api_key=_GROQ_KEY())
        with open(str(path), "rb") as f:
            result = client.audio.transcriptions.create(
                model="whisper-large-v3-turbo",
                file=f,
                language="es",
            )
        transcription = result.text
        if instruction and instruction.lower() not in ("transcribe", "transcribir"):
            return _gemini_text(f"{instruction}\n\nTranscripción:\n{transcription}")
        return transcription
    except Exception as e:
        return f"Error transcribiendo audio: {e}"


# ─── Data (CSV / Excel) ───────────────────────────────────────────────────────

def _process_data(path: Path, instruction: str) -> str:
    try:
        import pandas as pd
    except ImportError:
        return "pandas no instalado. Ejecuta: pip install pandas openpyxl"

    try:
        if path.suffix == ".csv":
            df = pd.read_csv(str(path))
        else:
            df = pd.read_excel(str(path))
    except Exception as e:
        return f"No pude leer el archivo de datos: {e}"

    rows, cols = df.shape
    preview = df.head(10).to_string(index=False)
    stats = df.describe(include="all").to_string()
    prompt = (
        f"{instruction}\n\n" if instruction else
        "Analiza estos datos en español. Describe qué contiene el dataset, "
        "estadísticas clave, valores atípicos y cualquier patrón interesante.\n\n"
    )
    prompt += f"Dataset: {rows} filas × {cols} columnas\nColumnas: {list(df.columns)}\n\nPrimeras filas:\n{preview}\n\nEstadísticas:\n{stats}"
    return _gemini_text(prompt[:14000])


# ─── Code ─────────────────────────────────────────────────────────────────────

def _process_code(path: Path, instruction: str) -> str:
    try:
        code = path.read_text(encoding="utf-8", errors="ignore")
    except Exception as e:
        return f"No pude leer el archivo: {e}"

    ext = path.suffix.lstrip(".")
    prompt = instruction or f"Analiza este código {ext} en español. Explica qué hace, identifica posibles problemas y sugiere mejoras:"
    return _gemini_text(f"{prompt}\n\n```{ext}\n{code[:12000]}\n```")


# ─── Documents (txt, md, docx) ───────────────────────────────────────────────

def _process_document(path: Path, instruction: str) -> str:
    text = ""
    if path.suffix == ".docx":
        try:
            from docx import Document
            doc = Document(str(path))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except ImportError:
            return "python-docx no instalado. Ejecuta: pip install python-docx"
        except Exception as e:
            return f"No pude leer el .docx: {e}"
    else:
        try:
            text = path.read_text(encoding="utf-8", errors="ignore")
        except Exception as e:
            return f"No pude leer el archivo: {e}"

    prompt = instruction or "Resume y analiza este documento en español, destacando los puntos más importantes:"
    return _gemini_text(f"{prompt}\n\n{text[:12000]}")


# ─── Gemini helpers ───────────────────────────────────────────────────────────

def _gemini_text(prompt: str) -> str:
    if not _GEMINI_KEY():
        return "GEMINI_API_KEY no configurada."
    try:
        import google.generativeai as genai
        genai.configure(api_key=_GEMINI_KEY())
        model = genai.GenerativeModel("gemini-2.0-flash")
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"Error de Gemini: {e}"


def _gemini_vision(path: Path, prompt: str) -> str:
    if not _GEMINI_KEY():
        return "GEMINI_API_KEY no configurada."
    try:
        import google.generativeai as genai
        from PIL import Image as PILImage
        genai.configure(api_key=_GEMINI_KEY())
        model = genai.GenerativeModel("gemini-2.0-flash")
        img = PILImage.open(str(path))
        response = model.generate_content([img, prompt])
        return response.text
    except Exception as e:
        return f"Error procesando imagen con Gemini: {e}"


def _fmt_bytes(size: int) -> str:
    for unit in ("B", "KB", "MB", "GB"):
        if size < 1024:
            return f"{size:.1f} {unit}"
        size /= 1024
    return f"{size:.1f} GB"
